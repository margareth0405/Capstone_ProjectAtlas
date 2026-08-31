"""In-memory text extraction for supported AI Detection documents."""

from pathlib import Path


class DocumentExtractionError(ValueError):
    """Raised when text cannot be safely extracted from an uploaded document."""


class DocumentTextExtractor:
    """Extract bounded plain text from PDF and modern Word documents."""

    supported_extensions = {".pdf", ".docx"}
    minimum_characters = 100
    maximum_characters = 20000

    def extract(self, uploaded_file):
        extension = Path(uploaded_file.name).suffix.lower()
        if extension == ".pdf":
            text = self._extract_pdf(uploaded_file)
        elif extension == ".docx":
            text = self._extract_docx(uploaded_file)
        else:
            raise DocumentExtractionError(
                "Upload a PDF or Word (.docx) document."
            )
        normalized = self._normalize(text)
        if len(normalized) < self.minimum_characters:
            raise DocumentExtractionError(
                "The document must contain at least 100 extractable characters. "
                "Scanned image-only PDFs require OCR before upload."
            )
        return normalized[: self.maximum_characters]

    @staticmethod
    def _extract_pdf(uploaded_file):
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise DocumentExtractionError(
                "PDF support is unavailable. Install the project requirements."
            ) from exc

        try:
            uploaded_file.seek(0)
            reader = PdfReader(uploaded_file)
            if reader.is_encrypted:
                raise DocumentExtractionError(
                    "Password-protected PDF files are not supported."
                )
            return chr(10).join(page.extract_text() or "" for page in reader.pages)
        except DocumentExtractionError:
            raise
        except Exception as exc:
            raise DocumentExtractionError(
                "ATLAS could not read this PDF file."
            ) from exc

    @staticmethod
    def _extract_docx(uploaded_file):
        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentExtractionError(
                "Word support is unavailable. Install the project requirements."
            ) from exc

        try:
            uploaded_file.seek(0)
            document = Document(uploaded_file)
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_text = [
                paragraph.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            ]
            return chr(10).join(paragraphs + table_text)
        except Exception as exc:
            raise DocumentExtractionError(
                "ATLAS could not read this Word document."
            ) from exc

    @staticmethod
    def _normalize(text):
        return chr(10).join(
            line.strip() for line in text.splitlines() if line.strip()
        )
